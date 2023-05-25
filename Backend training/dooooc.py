from docx import *
document =Document()
document.add_heading('Document Title',0)
p=document.add_paragraph("Nisar Ahmed", style="Intense Quote")
p.add_run("Nisar").bold=True


                    
document.save('demo.docx')
