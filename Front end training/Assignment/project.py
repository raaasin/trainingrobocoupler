from tkinter import *
from docx import *
top = Tk()
top.geometry("600x250")
top.title("Enquiry Form")
top.configure(bg="IndianRed3")

def magic():
    cusname=e1.get()
    proname=e2.get()
    email=e3.get()
    phone=e4.get()
    college=e5.get()
    pdesc=e6.get()
    datedeploy=e7.get()
    costpro=e8.get()
    d=Document()
    d.add_heading("Robocoupler Entry Form",0)

    d.add_paragraph("Customer Name: " + cusname)
    d.add_paragraph("Project Name: " + proname)
    d.add_paragraph("Email: " + email)
    d.add_paragraph("Phone Numer: " + phone)
    d.add_paragraph("College: " + college)
    d.add_paragraph("Project Description: " + pdesc)
    d.add_paragraph("Date of deployment: " + datedeploy)
    d.add_paragraph("Cost of the project: " + costpro)
    
    
    

    d.save('demo.docx')


head=Label(top,text="Robocoupler Enguiry Form",font=("Arial Bold",11),fg="White",bg="IndianRed3")
head.pack()

cname=Label(top,text="Customer Name:",fg="White",bg="IndianRed3")
cname.place(x=0,y=20)
e1=Entry(top)
e1.place(x=110,y=20)

pname=Label(top,text="Project Name:",fg="White",bg="IndianRed3")
pname.place(x=0,y=40)
e2=Entry(top)
e2.place(x=110,y=40)

mail=Label(top,text="Email:",fg="White",bg="IndianRed3")
mail.place(x=0,y=60)
e3=Entry(top)
e3.place(x=110,y=60)

pno=Label(top,text="Phone no:",fg="White",bg="IndianRed3")
pno.place(x=0,y=80)
e4=Entry(top)
e4.place(x=110,y=80)

csc=Label(top,text="College/School:",fg="White",bg="IndianRed3")
csc.place(x=0,y=100)
e5=Entry(top)
e5.place(x=110,y=100)

pd=Label(top,text="Project Description:",fg="White",bg="IndianRed3")
pd.place(x=0,y=120)
e6=Entry(top)
e6.place(x=110,y=120)

dd=Label(top,text="Date of Deploy:",fg="White",bg="IndianRed3")
dd.place(x=0,y=140)
e7=Entry(top)
e7.place(x=110,y=140)

cp=Label(top,text="Cost of project:",fg="White",bg="IndianRed3")
cp.place(x=0,y=160)
e8=Entry(top)
e8.place(x=110,y=160)

btn=Button(top,text="Submit",command = magic,bg="White",fg="IndianRed3",activebackground="IndianRed3")
btn.place(x=110,y=200)
top.mainloop()
